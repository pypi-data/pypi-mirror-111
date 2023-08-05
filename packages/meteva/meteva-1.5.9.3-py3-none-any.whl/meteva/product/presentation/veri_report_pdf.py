# -*- coding: utf-8 -*-
"""
Created on Sat Jun 19 11:58:32 2021

@author: 1
"""
import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from win32com import client
import os

doc = docx.Document()

path_root = r'H:\task\project\department\202004-检验系统研发\公报转pdf测试数据'

# 正文标题
folderList_0 = os.listdir(path_root)
folderName_0 = folderList_0[0]
title_0 = doc.add_paragraph()  # 正文大标题
run = title_0.add_run(folderName_0)  # 使用add_run添加文字
run.font.size = Pt(20)  # 设置字体大小
title_0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
run.bold = True  # 加粗
doc.styles['Normal'].font.name = '宋体'  # 设置字体
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# 添加标题
def AddHeadText(text, size, level):
    title_ = doc.add_heading('', level)
    # title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER# 标题居中
    title_run = title_.add_run(text)  # 添加标题内容
    title_run.font.size = Pt(size)  # 设置标题字体大小
    title_run.font.name = '宋体'  # 设置标题西文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置标题中文字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 字体颜色


def generateTitleCode(level):
    til = ''
    while level > 0:
        if len(til) < 1:
            til = str(level)
            print(til)
        else:
            til = str(level) + '.' + til
            print(til)
        level -= 1
    return til


folderList_1 = os.listdir(path_root + '/' + folderName_0)

tilCodeMap = {"1": 0}


def traverse(f, level, tilcode):
    fs = os.listdir(f)
    for f1 in fs:
        tmp_path = os.path.join(f, f1)
        if not os.path.isdir(tmp_path):
            if f1.endswith('.png'):
                pic = doc.add_picture(tmp_path, width=Inches(5))  # 添加图片
            if f1.endswith('.txt'):
                with open(tmp_path, encoding='utf-8', errors='ignore') as content_1:
                    data_1 = content_1.read()
                    p = doc.add_paragraph(data_1)
                    p.paragraph_format.first_line_indent = Cm(0.8)
                    p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
                    if data_1.startswith('图'):
                        # p = doc.add_paragraph(data_1)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            if tilcode not in tilCodeMap.keys():
                tilCodeMap[tilcode] = 0
            tilCodeMap[tilcode] = tilCodeMap[tilcode] + 1
            if tilcode == '1':
                AddHeadText(text=str(tilCodeMap[tilcode]) + '.' + f1[2:], size=12, level=level)
                traverse(tmp_path, level + 1, str(tilCodeMap[tilcode]) + '.')
            else:
                AddHeadText(text=tilcode + str(tilCodeMap[tilcode]) + '.' + f1[2:], size=12, level=level)
                traverse(tmp_path, level + 1, tilcode + str(tilCodeMap[tilcode]) + '.')


# 更新目录，前提是文档中已有目录
def update_doc(file):
    word = client.DispatchEx("Word.Application")  # 模拟打开 office

    try:

        doc = word.Documents.Open(file)  # 打开文件

        doc.TablesOfContents(1).Update()  # 更新目录

        doc.Close(SaveChanges=True)  # 关闭文档

        word.Quit()  # 退出

    except:

        print(file, "文件无目录！")


# 实现word转pdf
def doc2pdf(doc_name, pdf_name):
    """
    :word文件转pdf
    :param doc_name word文件名称
    :param pdf_name 转换后pdf文件名称
    """
    try:
        word = client.DispatchEx("Word.Application")
        if os.path.exists(pdf_name):
            os.remove(pdf_name)
        worddoc = word.Documents.Open(doc_name, ReadOnly=1)
        worddoc.SaveAs(pdf_name, FileFormat=17)
        worddoc.Close()
        return pdf_name
    except:
        return 0


if __name__ == "__main__":




    path = path_root + '/' + folderName_0
    traverse(path, 1, "1")

    doc_path = r'H:\task\project\department\202004-检验系统研发\进展\doc4.doc'
    doc.save(doc_path)
    update_doc(r'G:\Work\MODE\公报\python\doc4.doc')

    #rc = doc2pdf(doc_path, output_pdf)



